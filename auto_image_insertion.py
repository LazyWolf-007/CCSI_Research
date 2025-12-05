"""
Smart CCSI Appendix Generator - Automatically finds your images by pattern matching
No need to rename files!
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import glob

def find_image(folder, pattern_keywords):
    """
    Smart search for images using keywords
    Returns the first matching file found
    """
    # Try exact patterns first
    for keyword in pattern_keywords:
        # Try with .png
        matches = glob.glob(os.path.join(folder, f"*{keyword}*.png"), recursive=False)
        if matches:
            return matches[0]
        # Try with .jpg
        matches = glob.glob(os.path.join(folder, f"*{keyword}*.jpg"), recursive=False)
        if matches:
            return matches[0]
        # Try with .jpeg
        matches = glob.glob(os.path.join(folder, f"*{keyword}*.jpeg"), recursive=False)
        if matches:
            return matches[0]
    return None

def create_ccsi_appendix_smart(image_folder_path, output_docx_path):
    """
    Creates CCSI Appendix - automatically finds images by name pattern
    """
    
    # Validate paths
    if not os.path.exists(image_folder_path):
        print(f"❌ ERROR: Folder not found: {image_folder_path}")
        return
    
    if not output_docx_path.endswith('.docx'):
        output_docx_path = output_docx_path + '.docx'
    
    # Create output directory if needed
    output_dir = os.path.dirname(output_docx_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print(f"📁 Looking for images in: {image_folder_path}")
    print(f"💾 Will save to: {output_docx_path}\n")
    
    # Create document
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Title Page
    title = doc.add_heading('GRAPHICAL APPENDIX', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Full Visual Dataset for the CCSI Research Paper')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    main_title = doc.add_heading(
        '"Cultural Cohesion and Social Inclusivity Index (CCSI): '
        'A Data-Driven Study of India\'s Civilizational Patterns"',
        level=1
    )
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author = doc.add_paragraph('Author: Tejas Pradip Pawar')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.runs[0].bold = True
    
    # Intro
    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(6)
    intro.paragraph_format.space_after = Pt(6)
    intro.paragraph_format.line_spacing = 1.0
    intro.add_run(
        'This appendix contains all 17 research figures used during analysis. '
        'Due to IEEE page-limit constraints, only 3 visualizations were placed in the main paper. '
        'This appendix provides the complete dataset-driven graphical evidence, including '
        'historical CCSI timelines, dimensional heatmaps, and ML-based clustering diagnostics.'
    )
    
    doc.add_page_break()
    
    # Define figures with smart search keywords
    figures = [
        ('SECTION A — CCSI Time Series (Historical)', [
            ('1️⃣ CCSI over Time – All Regions', ['all_regions_comparison', 'all_regions', 'comparison']),
            ('2️⃣ 3-Era Rolling Mean of CCSI – All Regions', ['rolling_mean', 'rolling', '3_era']),
            ('3️⃣ CCSI Evolution – Unified Timeline Plot', ['timeline_all', 'evolution', 'unified']),
        ]),
        
        ('SECTION B — Region-Specific CCSI Curves', [
            ('4️⃣ CCSI over Time – Deccan', ['timeline_Deccan', 'Deccan_timeline', 'deccan']),
            ('5️⃣ CCSI over Time – Gangetic North', ['Gangetic_North', 'gangetic', 'north_timeline']),
            ('6️⃣ CCSI over Time – Northeast', ['timeline_Northeast', 'northeast_timeline', 'northeast']),
            ('7️⃣ CCSI over Time – Northwest (Punjab–Gandhāra)', ['Northwest', 'Punjab', 'Gandhara', 'northwest_punjab']),
            ('8️⃣ CCSI over Time – Tamilakam', ['timeline_Tamilakam', 'tamilakam_timeline', 'tamilakam']),
        ]),
        
        ('SECTION C — D1–D5 Dimension Heatmaps', [
            ('9️⃣ D1-D5 Dimensional Heatmap – Deccan', ['heatmap_Deccan', 'dimensions_deccan', 'deccan_heat']),
            ('🔟 D1-D5 Dimensional Heatmap – Gangetic North', ['heatmap_Gangetic', 'dimensions_gangetic', 'gangetic_heat']),
            ('1️⃣1️⃣ D1-D5 Dimensional Heatmap – Northeast', ['heatmap_Northeast', 'dimensions_northeast', 'northeast_heat']),
            ('1️⃣2️⃣ D1-D5 Dimensional Heatmap – Northwest (Punjab–Gandhāra)', ['heatmap_Northwest', 'dimensions_northwest', 'northwest_heat']),
            ('1️⃣3️⃣ D1-D5 Dimensional Heatmap – Tamilakam', ['heatmap_Tamilakam', 'dimensions_tamilakam', 'tamilakam_heat']),
        ]),
        
        ('SECTION D — Machine Learning Outputs', [
            ('1️⃣4️⃣ K-Means Clusters (k=4) on PCA Components (D1–D5 totals)', ['kmeans', 'clusters', 'k_means']),
            ('1️⃣5️⃣ PCA of D1–D5 Totals – All Eras & Regions', ['pca_d1_d5', 'pca_scatter', 'pca_totals']),
            ('1️⃣6️⃣ PCA: K-Means with Modern Eras Highlighted', ['modern_highlight', 'clusters_modern', 'pca_modern']),
            ('1️⃣7️⃣ PCA: Historical vs Modern Eras Separation', ['E_vs_M', 'historical_modern', 'pca_historical']),
        ]),
    ]
    
    found_count = 0
    missing_count = 0
    
    # Insert each section with figures
    for section_title, section_figures in figures:
        # Section header
        header = doc.add_heading(f'🔹 {section_title}', level=2)
        header.paragraph_format.space_before = Pt(10)
        header.paragraph_format.space_after = Pt(6)
        
        # Add each figure
        for fig_title, search_keywords in section_figures:
            # Figure title
            fig_para = doc.add_paragraph()
            fig_para.add_run(fig_title).bold = True
            fig_para.paragraph_format.space_before = Pt(4)
            fig_para.paragraph_format.space_after = Pt(2)
            
            # Smart search for image
            image_path = find_image(image_folder_path, search_keywords)
            
            if image_path:
                # Insert image
                doc.add_picture(image_path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                last_paragraph.paragraph_format.space_after = Pt(6)
                last_paragraph.paragraph_format.space_before = Pt(0)
                
                filename = os.path.basename(image_path)
                print(f"✅ Found: {filename}")
                found_count += 1
            else:
                # Placeholder
                placeholder = doc.add_paragraph(f'[Image not found - searched for: {", ".join(search_keywords)}]')
                placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                placeholder.runs[0].italic = True
                placeholder.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                placeholder.paragraph_format.space_after = Pt(6)
                print(f"❌ Missing: {fig_title}")
                missing_count += 1
    
    # Closing
    doc.add_page_break()
    closing_header = doc.add_heading('Closing Statement', level=2)
    closing_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    closing_text = doc.add_paragraph(
        'This appendix is supplementary material accompanying the IEEE paper submission. '
        'It is not intended for print publication but for peer-review evaluation of dataset '
        'completeness, methodological transparency, and reproducibility of results.\n\n'
        'All figures were generated programmatically using Python (pandas, numpy, matplotlib, '
        'seaborn, sklearn). Full code and dataset are available upon request.'
    )
    
    end_mark = doc.add_paragraph('— End of Graphical Appendix —')
    end_mark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_mark.runs[0].italic = True
    
    # Save
    doc.save(output_docx_path)
    
    print(f"\n{'='*60}")
    print(f"✅ Document created: {output_docx_path}")
    print(f"📊 Images found: {found_count}/17")
    print(f"❌ Images missing: {missing_count}/17")
    if missing_count > 0:
        print(f"\n💡 Tip: Check the folder for remaining images")
    print(f"{'='*60}")


# === USAGE ===
if __name__ == "__main__":
    
    # CHANGE THESE TWO LINES TO YOUR ACTUAL PATHS!
    IMAGE_FOLDER = r"C:\Users\tejas\Civilizational_Strength_Index_ResearchPaper\figures"
    OUTPUT_FILE = r"C:\Users\tejas\Civilizational_Strength_Index_ResearchPaper\CCSI_Appendix_COMPACT.docx"
    
    print("🚀 Creating CCSI Graphical Appendix with smart image detection...\n")
    create_ccsi_appendix_smart(IMAGE_FOLDER, OUTPUT_FILE)